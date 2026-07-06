import os
import json
import time
import httpx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv('.env')
GROQ_API_KEY = os.environ.get('GROQ_API_KEY')

if not GROQ_API_KEY:
    print("GROQ_API_KEY not found in .env")
    exit(1)

context = """
PROJECT AURA FASHION AI - SUMMARY FOR DOC GENERATION
Architecture: Flutter Mobile App -> FastAPI Backend -> SQLite/Qdrant/Redis -> LangGraph Agents -> Groq LLM + Mediapipe + Stable Diffusion.
Current State: Uses pre-trained models via APIs. Features zero-card free stack, voice STT/TTS pipelines, and RAG based styling.
PROPOSED FINE-TUNING ARCHITECTURE: The roadmap proposes upgrading from zero-shot prompting to a fine-tuned reasoning model in 3 phases:
1. Principles: Teach the model fashion geometry, color theory, and fabric science (~3k examples).
2. Reasoning: Teach step-by-step reasoning via <think> blocks factoring body, skin, climate, budget, culture (~8k examples).
3. Novel Synthesis: Combine extreme constraints for advanced generalization (~4k examples).
You must write documentation describing the existing systems AND this proposed fine-tuning evolution.
"""

sections = [
    ("1. Declaration", "Standard university major project declaration stating the original work of the project. Write ~200 words."),
    ("2. Acknowledgement", "Standard acknowledgement thanking mentors, guides, and the open-source community. Write ~200 words."),
    ("3. Abstract", "A 400-word abstract summarizing Aura, an AI-powered personal fashion designer, its current capabilities (RAG, Groq LLM, Mediapipe body analysis, Flutter app), and the proposed upgrade to a reasoning-first fine-tuned model."),
    ("4. Table of Contents", "Just a placeholder page saying 'TABLE OF CONTENTS - Please generate automatically in Word.'"),
    ("5. List of Figures", "Just a placeholder page saying 'LIST OF FIGURES - Please generate automatically in Word.'"),
    ("6. List of Tables", "Just a placeholder page saying 'LIST OF TABLES - Please generate automatically in Word.'"),
    ("7. List of Abbreviations", "A comprehensive list of abbreviations used in the project (e.g., RAG, LLM, VLM, STT, TTS, API, JWT). Detail at least 15 abbreviations."),
    ("8. Introduction", "Detailed introduction: Background of fashion AI, Problem Statement (generic fashion advice vs personalized reasoning), Objectives, Project Scope, Motivation, and Overview. Write ~800 words."),
    ("9. Literature Review", "Detailed literature review comparing industry solutions, research papers on fashion recommendation, current AI technologies, and highlighting the novelty of Aura's reasoning-first approach. Write ~800 words."),
    ("10. Requirement Analysis", "Functional Requirements, Non-functional Requirements, System Constraints, Assumptions, and User Roles. Use lists and detailed descriptions. Write ~600 words."),
    ("11. System Architecture", "Overall Architecture, Technology Stack, Folder Structure, Component Description, Communication Flow, Request Flow, Response Flow, Data Flow, Deployment Architecture, Infrastructure, Cloud Architecture. Include a MERMAID diagram for Overall Architecture. Write ~800 words."),
    ("12. Module Description", "Explain EVERY major module: Flutter App, FastAPI Backend, LangGraph Agents (Master, Body Analyzer, Fashion RAG, Stylist, Tailoring Calculator, Product Retrieval), Voice Pipeline, and Image Generation. Detailed descriptions of each. Write ~1000 words."),
    ("13. Database Design", "ER Diagram (MERMAID format), Collections/Tables (Users, Sessions, Wardrobe), Relationships, Indexes, Storage Strategy (SQLite + Qdrant + Redis). Write ~600 words."),
    ("14. AI/ML Architecture", "Current pretrained architecture (Groq Llama 3.1, Mediapipe, Stable Diffusion) AND the PROPOSED Fine-tuned architecture based on the finetuning guide (3 phases: Principles, Reasoning, Novel Synthesis). Detail the training pipeline, inference pipeline, evaluation, prompt engineering, and model lifecycle. Include a MERMAID diagram for the training pipeline. Write ~1000 words."),
    ("15. Algorithms", "Explain every important algorithm (e.g., Body Shape Classification using WHR/SHR, RAG Similarity Search, Cost-Per-Wear calculation, Stylist Chain-of-Thought). Provide pseudo algorithms and complexity analysis. Write ~600 words."),
    ("16. Workflow", "End-to-end execution flow from user voice input, STT, Agent processing, Body Analysis, RAG, Stylist Reasoning, to Image Gen and TTS output. Include a MERMAID Flowchart. Write ~600 words."),
    ("17. Sequence Diagrams", "Provide detailed sequence diagrams in MERMAID format for: User Login, Chat/Voice Query Processing, and Design Generation. Include detailed descriptive text for each step. Write ~600 words."),
    ("18. Activity Diagrams", "Provide an Activity Diagram in MERMAID format for the user journey (capturing body mesh -> getting outfit -> saving to wardrobe). Include descriptive text. Write ~500 words."),
    ("19. Class Diagrams", "Provide a Class Diagram in MERMAID format for the backend data models and agent structure. Include descriptive text. Write ~500 words."),
    ("20. Component Diagrams", "Provide a Component Diagram in MERMAID format showing the physical and logical components. Include descriptive text. Write ~500 words."),
    ("21. Deployment Diagram", "Provide a Deployment Diagram in MERMAID format showing Render, Docker, SQLite, and Mobile App. Include descriptive text. Write ~500 words."),
    ("22. API Documentation", "Detailed API Documentation for: /auth, /chat/message, /session/design-flow, /voice/stream, /wardrobe/save. Detail Inputs, Outputs, Auth, Error handling. Write ~800 words."),
    ("23. Security", "Details on Authentication (JWT), Authorization, Secrets management, Privacy (zero-card, local DB), Rate limiting (SlowAPI), and Validation. Write ~600 words."),
    ("24. Performance", "Details on Caching (Redis), Optimization, Latency management, Scalability (Stateless FastAPI), GPU utilization (HF Inference fallback). Write ~600 words."),
    ("25. Testing", "Unit testing, Integration testing, System testing, AI evaluation metrics, and Acceptance testing. (17 existing tests in pytest). Write ~600 words."),
    ("26. Results", "Discuss current implementation results vs expected improvements after fine-tuning (latency vs accuracy trade-offs, personalization quality). Include a plain text table of Performance Comparison. Write ~600 words."),
    ("27. Future Enhancements", "Extensive details on future enhancements, strongly referencing continuous learning, expanding the finetuning dataset, handling new fashion trends, and scaling the vector DB. Write ~600 words."),
    ("28. Conclusion", "A strong, academic conclusion summarizing the project's achievements and the impact of the proposed fine-tuning architecture. Write ~400 words."),
    ("29. References", "List of references in IEEE format (simulated references to Groq, Meta Llama, LangChain, Flutter, FastAPI, and fashion science literature). Write ~200 words.")
]

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

title = doc.add_heading('Aura Fashion AI', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Comprehensive Project Documentation & Architecture Evolution')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

def generate_content(section_title, instructions):
    prompt = f"""
You are a Principal Software Architect. Write a highly detailed section for a Major Project Report.
Project: Aura Fashion AI (multilingual voice AI fashion stylist).

CONTEXT:
{context}

SECTION TO WRITE:
{section_title}

INSTRUCTIONS:
{instructions}

FORMATTING RULES:
- Do NOT use markdown symbols like **, *, # for formatting.
- For headings, start the line with 'HEADING1: ' or 'HEADING2: ' or 'HEADING3: '.
- For bullet points, start the line with '- '.
- For tables, use plain text with columns separated by | (e.g., Col1 | Col2 | Col3).
- For Mermaid diagrams, start the block with a line exactly reading 'MERMAID_START' and end with 'MERMAID_END'.
- Write to meet the word count specified. Be extremely exhaustive, technical, and academic.
"""
    
    # Prompt is ~350 tokens. max_tokens=2000. Total requested = 2350 tokens.
    # 6000 TPM limit / 2350 = ~2.5 requests per minute. 
    # Sleep 25 seconds between requests to guarantee 2.4 requests/min, staying under 6000 TPM.

    for attempt in range(5):
        try:
            r = httpx.post("https://api.groq.com/openai/v1/chat/completions", headers={
                "Authorization": f"Bearer {GROQ_API_KEY}"
            }, json={
                "model": "llama-3.1-8b-instant",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 2000
            }, timeout=60.0)
            
            if r.status_code == 200:
                return r.json()["choices"][0]["message"]["content"]
            elif r.status_code in [429, 413]:
                # Rate limit hit, backoff heavily
                print(f"Rate limited on {section_title}, sleeping 40s...")
                time.sleep(40)
            else:
                print(f"Error {r.status_code}: {r.text}")
                time.sleep(10)
        except Exception as e:
            print(f"Exception on {section_title}: {e}")
            time.sleep(10)
    return "Failed to generate content due to strict API limits."

for title, instructions in sections:
    print(f"Generating {title}...")
    doc.add_heading(title, level=1)
    
    content = generate_content(title, instructions)
    
    in_mermaid = False
    mermaid_content = []
    
    for line in content.split('\n'):
        line_stripped = line.strip()
        
        if line_stripped == 'MERMAID_START':
            in_mermaid = True
            continue
        elif line_stripped == 'MERMAID_END':
            in_mermaid = False
            p = doc.add_paragraph()
            p.add_run("```mermaid\n" + "\n".join(mermaid_content) + "\n```").font.name = 'Courier New'
            mermaid_content = []
            continue
            
        if in_mermaid:
            mermaid_content.append(line)
            continue
            
        if line.startswith('HEADING1:'):
            doc.add_heading(line.replace('HEADING1:', '').strip(), level=2)
        elif line.startswith('HEADING2:'):
            doc.add_heading(line.replace('HEADING2:', '').strip(), level=3)
        elif line.startswith('HEADING3:'):
            doc.add_heading(line.replace('HEADING3:', '').strip(), level=4)
        elif line.startswith('- '):
            p = doc.add_paragraph(line.replace('- ', '', 1).strip(), style='List Bullet')
        elif '|' in line and len(line.split('|')) > 1:
            p = doc.add_paragraph(line)
            p.runs[0].font.name = 'Courier New'
        elif line_stripped != "":
            doc.add_paragraph(line_stripped)
            
    doc.add_page_break()
    time.sleep(25) # CRITICAL: Stay under 6000 TPM limit

output_path = r"C:\Users\vishn\Downloads\Aura_Gem_Complete_Project_Documentation.docx"
doc.save(output_path)
print(f"[SUCCESS] Document successfully saved to {output_path}")
